from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 建立 Word 文件
doc = Document()

# 設定預設字型為 Noto Sans TC（需電腦已安裝）
style = doc.styles['Normal']
style.font.name = 'Noto Sans TC'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans TC')
style.font.size = Pt(12)

# 封面
doc.add_paragraph('\n'*10)  # 上方間距
title = doc.add_paragraph("服務學習活動方案")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(24)
title.runs[0].font.bold = True
title.runs[0].font.color.rgb = RGBColor(46, 64, 83)

doc.add_paragraph('\n'*5)
subtitle = doc.add_paragraph("申請職位：課外活動行政專員\n應徵者：XXX")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

doc.add_page_break()

# 章節內容
sections = [
    ("🌱 活動方案名稱", "社區健康小幫手：大學生走入社區的健康促進服務學習計畫"),
    ("🎯 活動目標", "1. 結合醫學、護理、生命科學專業，讓學生將所學應用於真實場域。\n"
                     "2. 協助社區民眾提升健康知識與自我照護能力。\n"
                     "3. 培養學生公共服務態度、團隊合作與跨域溝通能力。\n"
                     "4. 建立大學與社區的長期合作平台，實踐大學社會責任（USR）。"),
    ("👥 參與對象", "‧ 陽明交通大學在學學生（醫學、護理、藥學、公共衛生、生命科學系為主，亦開放跨院系參與）。\n"
                     "‧ 服務對象：鄰近社區居民、銀髮族、弱勢家庭。"),
    ("📅 活動內容與流程", "1️⃣ 前期培訓（2週）\n"
                     "- 工作坊：健康促進與社區服務倫理\n"
                     "- 跨系小組分工（健康檢測組、衛教宣導組、生活關懷組）\n"
                     "- 學習基本服務技巧（血壓量測、健康諮詢、簡易心肺復甦術、飲食運動衛教）\n\n"
                     "2️⃣ 社區服務（4週）\n"
                     "- 每週一次，深入合作社區據點\n"
                     "- 活動：健康檢測、衛教小講堂、生活關懷、親子互動\n\n"
                     "3️⃣ 成果回饋（1週）\n"
                     "- 成果分享會，學生小組提交服務反思報告"),
    ("📊 預期效益", "‧ 對學生：提升公共服務能力、加深專業知識實踐、培養同理心\n"
                     "‧ 對社區：居民獲得基礎健康知識、增強自我健康管理能力\n"
                     "‧ 對學校：展現大學社會責任，強化校園與社區連結"),
    ("💡 評估方式", "1. 學生：服務反思報告 + 同儕互評 + 導師回饋\n"
                     "2. 社區：問卷調查民眾滿意度\n"
                     "3. 學校：活動成果納入USR案例")
]

for title, text in sections:
    # 標題
    p = doc.add_paragraph(title)
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = RGBColor(26, 82, 118)
    
    # 分隔線
    tbl = doc.add_table(rows=1, cols=1)
    tbl_cell = tbl.cell(0, 0)
    tbl_cell.width = Cm(16)
    tbl_cell.height = Cm(0.3)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), "D6EAF8")  # 淺藍色
    tbl_cell._tc.get_or_add_tcPr().append(shading_elm)
    tbl_cell.text = ""
    
    # 內文
    for line in text.split("\n"):
        para = doc.add_paragraph(line, style='Normal')
        para.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("\n")  # 段落間距

# 儲存 Word 文件
doc.save("服務學習活動方案.docx")
print("已生成 Word 文件: 服務學習活動方案.docx")